#!/usr/bin/env python3
"""generate_submission_report.py — .docx Submission Report Generator.

For each candidate of a given job that has a t2_score (or t1_score) >= PASS_THRESHOLD
and does NOT yet have a `traffic_rpt` attachment, this script:

  1. Fetches the Airtable record (ai_detailed_json, ai_summary, full_name, job_name, cv_text)
  2. Loads the rubric from Airtable (for category classification: technical vs soft_skill)
  3. Infers candidate metadata (location, nationality, relevant experience) via LLM
  4. Builds a .docx entirely in Python using python-docx (no Node.js required)
  5. Uploads the .docx to the `traffic_rpt` attachment field in Airtable

Usage:
  python generate_submission_report.py <job_id>
"""

from __future__ import annotations

import io
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openai import OpenAI

from airtable_client import AirtableClient
from config import Config
from utils import safe_filename


# ══════════════════════════════════════════════════════════════════════════════
# Colour palette (matches build_report.js)
# ══════════════════════════════════════════════════════════════════════════════
_C = {
    "darkBlue":  (0x1F, 0x4E, 0x79),
    "midBlue":   (0x2E, 0x75, 0xB6),
    "lightBlue": (0xEB, 0xF3, 0xFB),
    "headerBg":  (0xD6, 0xE4, 0xF0),
    "green":     (0x27, 0xAE, 0x60),
    "amber":     (0xE6, 0x7E, 0x22),
    "red":       (0xC0, 0x39, 0x2B),
    "gray":      (0x59, 0x59, 0x59),
    "white":     (0xFF, 0xFF, 0xFF),
    "black":     (0x00, 0x00, 0x00),
}

def _rgb(key: str) -> RGBColor:
    t = _C[key]
    return RGBColor(*t)

def _hex(key: str) -> str:
    t = _C[key]
    return "%02X%02X%02X" % t


# ══════════════════════════════════════════════════════════════════════════════
# python-docx XML helpers
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:shd")):
        tcPr.remove(s)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(
    cell,
    top: Optional[Tuple] = None,
    right: Optional[Tuple] = None,
    bottom: Optional[Tuple] = None,
    left: Optional[Tuple] = None,
) -> None:
    """Each border is None (transparent) or a (style, size, hex_color) tuple."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(s)
    tcBorders = OxmlElement("w:tcBorders")
    for side, bdef in [("top", top), ("right", right), ("bottom", bottom), ("left", left)]:
        el = OxmlElement(f"w:{side}")
        if bdef is None:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:color"), "FFFFFF")
            el.set(qn("w:space"), "0")
        else:
            style, sz, color = bdef
            el.set(qn("w:val"), style)
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), color)
            el.set(qn("w:space"), "0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=100, bottom=100, left=140, right=140) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(s)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:tcW")):
        tcPr.remove(s)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.insert(0, tcW)


def _set_table_width(table, width_dxa: int) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for s in tblPr.findall(qn("w:tblW")):
        tblPr.remove(s)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_table_fixed(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _clear_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for s in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(s)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "FFFFFF")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_para_spacing(para, before: int = 0, after: int = 0) -> None:
    """before/after in twips (1/20 pt)."""
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn("w:spacing")):
        pPr.remove(s)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _set_para_border(para, side: str, color: str, size: int) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:color"), color)
    el.set(qn("w:space"), "1")
    pBdr.append(el)


def _clear_cell(cell) -> None:
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def _add_run(
    para,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size_pt: float = 10.0,
    color: str = "black",
    all_caps: bool = False,
):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = _rgb(color)
    if all_caps:
        run.font.all_caps = True
    return run


def _add_emoji_run(para, text: str, size_pt: float = 12.0):
    """Add an emoji run using Segoe UI Emoji so coloured rendering works.

    Characters like ⚠ (U+26A0) default to *text* presentation and only become
    coloured emoji when the variation selector U+FE0F is present AND the run is
    rendered with an emoji-capable font.  Without an explicit font Word falls
    back to Calibri which ignores U+FE0F, giving a monochrome glyph.  Setting
    all four rFonts slots to Segoe UI Emoji (the standard Windows emoji font)
    fixes this for every emoji used in the report.
    """
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    # Inject <w:rFonts> for all four slots so Word uses Segoe UI Emoji
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Segoe UI Emoji")
    rPr.insert(0, rFonts)
    return run


# ══════════════════════════════════════════════════════════════════════════════
# .docx builder (pure Python)
# ══════════════════════════════════════════════════════════════════════════════

def build_docx(payload: Dict) -> bytes:
    """Build the candidate submission .docx from a payload dict."""

    full_name           = payload.get("full_name", "")
    job_name            = payload.get("job_name", "")
    overall_score       = float(payload.get("overall_score") or 0)
    recommendation      = payload.get("recommendation", "FAIL")
    ai_summary          = payload.get("ai_summary", "")
    compliance          = payload.get("compliance") or []
    technical           = payload.get("technical") or []
    soft_skill          = payload.get("soft_skill") or []
    nice_to_have        = payload.get("nice_to_have") or []
    location            = payload.get("location", "[Location — please verify]")
    nationality         = payload.get("nationality", "[Nationality — please verify]")
    relevant_experience = payload.get("relevant_experience", "[Experience — please verify]")
    report_date         = payload.get("report_date", "")

    rec_color = "green" if recommendation == "PASS" else ("amber" if recommendation == "REVIEW" else "red")

    # Border presets (style, size_eighths_pt, hex_color)
    OUTER  = ("single", 4, "BBBBBB")
    INNER  = ("single", 4, "D1D1D1")
    NONE_B = ("none",   0, "FFFFFF")

    # Column widths in DXA (1 DXA = 1/1440 inch). Total = 9360 ≈ 6.5 in (A4 with margins)
    W_REQ, W_FIT, W_EV = 3680, 1000, 4680

    def score_to_fit(score) -> Tuple[str, str, str]:
        """Returns (emoji, label, colour_key)."""
        s = float(score) if score is not None else 0
        if s >= 4:  return ("\u2705",           "Meets",   "green")  # ✅
        if s >= 2:  return ("\u26A0\uFE0F",    "Partial", "amber")  # ⚠️
        return          ("\u274C",           "Not Met", "red")    # ❌

    def compliance_to_fit(status) -> Tuple[str, str, str]:
        if (status or "").upper() == "PASS":
            return ("\u2705", "Meets",   "green")  # ✅
        return     ("\u274C", "Not Met", "red")    # ❌

    # ── Document + page setup ─────────────────────────────────────────────────
    doc = Document()
    # Remove default Normal style spacing
    ns = doc.styles["Normal"]
    ns.font.name = "Arial"
    ns.font.size = Pt(10)
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after  = Pt(0)

    sec = doc.sections[0]
    sec.page_width    = Inches(8.268)    # A4
    sec.page_height   = Inches(11.693)
    sec.top_margin    = Inches(1224 / 1440)
    sec.right_margin  = Inches(1080 / 1440)
    sec.bottom_margin = Inches(1080 / 1440)
    sec.left_margin   = Inches(1080 / 1440)

    # Remove the one default empty paragraph that Document() creates
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)

    # ── Header info ───────────────────────────────────────────────────────────
    # Name
    p = doc.add_paragraph()
    _set_para_spacing(p, 0, 60)
    _add_run(p, full_name, bold=True, size_pt=16, color="darkBlue")

    # Role + Experience
    p = doc.add_paragraph()
    _set_para_spacing(p, 0, 40)
    _add_run(p, "Role Applied: ",                   bold=True, size_pt=12)
    _add_run(p, job_name,                           bold=True, size_pt=12, color="darkBlue")
    _add_run(p, "     |     Relevant Experience: ", bold=True, size_pt=12)
    _add_run(p, relevant_experience,                           size_pt=12)

    # Location + Notice Period
    p = doc.add_paragraph()
    _set_para_spacing(p, 0, 40)
    _add_run(p, "Location: ",                       bold=True, size_pt=11)
    _add_run(p, location,                                      size_pt=11)
    _add_run(p, "     |     Notice Period: ",        bold=True, size_pt=11)
    _add_run(p, "[To be filled by recruiter]",                 size_pt=11, color="gray")

    # Nationality
    p = doc.add_paragraph()
    _set_para_spacing(p, 0, 40)
    _add_run(p, "Nationality: ", bold=True, size_pt=11)
    _add_run(p, nationality,                size_pt=11)

    # Score + Recommendation
    p = doc.add_paragraph()
    _set_para_spacing(p, 0, 80)
    _add_run(p, f"AI Score: {int(overall_score)}/100     ", bold=True, size_pt=11, color="gray")
    _add_run(p, f"  {recommendation}  ",                   bold=True, size_pt=11, color=rec_color)
    _add_run(p, f"     Report Date: {report_date}",                    size_pt=10, color="gray")

    # ── Local helpers using the open doc ─────────────────────────────────────
    def add_section_label(text: str) -> None:
        p = doc.add_paragraph()
        _set_para_spacing(p, 240, 100)
        _set_para_border(p, "bottom", _hex("midBlue"), 8)
        _add_run(p, text, bold=True, size_pt=11, color="darkBlue", all_caps=True)

    def add_gap(pt: float = 6) -> None:
        p = doc.add_paragraph()
        _set_para_spacing(p, 0, int(pt * 20))
        _add_run(p, "", size_pt=pt)

    def add_assessment_box(text: str) -> None:
        tbl = doc.add_table(rows=1, cols=1)
        _set_table_width(tbl, 9360)
        _clear_table_borders(tbl)
        _set_table_fixed(tbl)
        cell = tbl.rows[0].cells[0]
        _set_cell_width(cell, 9360)
        _set_cell_bg(cell, _hex("lightBlue"))
        _set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
        _set_cell_borders(cell,
            top    = ("single", 6, _hex("midBlue")),
            right  = ("single", 6, _hex("midBlue")),
            bottom = ("single", 6, _hex("midBlue")),
            left   = ("single", 6, _hex("midBlue")),
        )
        _clear_cell(cell)
        p = cell.add_paragraph()
        _set_para_spacing(p, 0, 0)
        _add_run(p, text, size_pt=10.5)

    def add_req_table(items: List[Dict], req_key="requirement", ev_key="evidence") -> None:
        if not items:
            return
        tbl = doc.add_table(rows=1 + len(items), cols=3)
        _set_table_width(tbl, W_REQ + W_FIT + W_EV)
        _clear_table_borders(tbl)
        _set_table_fixed(tbl)

        # Header row
        hdr_row = tbl.rows[0]
        for col_i, (cell, label) in enumerate(zip(hdr_row.cells, ["Requirement", "Fit", "Evidence"])):
            is_left  = col_i == 0
            is_right = col_i == 2
            _set_cell_width(cell, [W_REQ, W_FIT, W_EV][col_i])
            _set_cell_bg(cell, _hex("headerBg"))
            _set_cell_margins(cell)
            _set_cell_borders(cell,
                top    = OUTER,
                bottom = INNER,
                left   = OUTER if is_left  else NONE_B,
                right  = OUTER if is_right else NONE_B,
            )
            _clear_cell(cell)
            p = cell.add_paragraph()
            _set_para_spacing(p, 0, 0)
            _add_run(p, label, bold=True, size_pt=10.5, color="darkBlue")

        # Data rows
        last_i = len(items) - 1
        for row_i, item in enumerate(items):
            is_last = row_i == last_i
            data_row = tbl.rows[row_i + 1]

            if "score" in item:
                icon, fit_label, fit_color = score_to_fit(item.get("score"))
            else:
                icon, fit_label, fit_color = compliance_to_fit(item.get("status", "FAIL"))

            req_text = item.get(req_key) or item.get("skill") or ""
            ev_text  = item.get(ev_key)  or "Not demonstrated"

            def borders(is_left_edge: bool, is_right_edge: bool) -> Dict:
                return {
                    "top":    INNER,
                    "bottom": OUTER if is_last else INNER,
                    "left":   OUTER if (is_last and is_left_edge)  else INNER,
                    "right":  OUTER if (is_last and is_right_edge) else INNER,
                }

            # Requirement
            c = data_row.cells[0]
            _set_cell_width(c, W_REQ)
            _set_cell_bg(c, _hex("white"))
            _set_cell_margins(c)
            b = borders(True, False)
            _set_cell_borders(c, top=b["top"], bottom=b["bottom"], left=b["left"], right=b["right"])
            _clear_cell(c)
            p = c.add_paragraph(); _set_para_spacing(p, 0, 0)
            _add_run(p, req_text, size_pt=10)

            # Fit
            c = data_row.cells[1]
            _set_cell_width(c, W_FIT)
            _set_cell_bg(c, _hex("white"))
            _set_cell_margins(c, top=100, bottom=100, left=100, right=100)
            b = borders(False, False)
            _set_cell_borders(c, top=b["top"], bottom=b["bottom"], left=b["left"], right=b["right"])
            _clear_cell(c)
            p = c.add_paragraph(); _set_para_spacing(p, 0, 0)
            _add_emoji_run(p, icon + " ", size_pt=11)
            _add_run(p, fit_label, size_pt=10, color=fit_color)

            # Evidence
            c = data_row.cells[2]
            _set_cell_width(c, W_EV)
            _set_cell_bg(c, _hex("white"))
            _set_cell_margins(c)
            b = borders(False, True)
            _set_cell_borders(c, top=b["top"], bottom=b["bottom"], left=b["left"], right=b["right"])
            _clear_cell(c)
            p = c.add_paragraph(); _set_para_spacing(p, 0, 0)
            _add_run(p, ev_text, size_pt=10)

    # ── Body sections ─────────────────────────────────────────────────────────
    add_section_label("Recruiter's Assessment")
    add_gap(6)
    add_assessment_box(ai_summary)
    add_gap(10)

    if compliance:
        add_section_label("Compliance Gates")
        add_gap(6)
        add_req_table(compliance, "requirement", "evidence")
        add_gap(12)

    if technical:
        add_section_label("Must-Have Requirements")
        add_gap(6)
        add_req_table(technical, "requirement", "evidence")
        add_gap(12)

    if soft_skill:
        add_section_label("Soft-Skill Indicators")
        add_gap(6)
        add_req_table(soft_skill, "requirement", "evidence")
        add_gap(12)

    if nice_to_have:
        add_section_label("Nice-to-Have")
        add_gap(6)
        add_req_table(nice_to_have, "skill", "evidence")
        add_gap(16)

    # ── Footer ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    _set_para_spacing(p, 240, 0)
    _set_para_border(p, "top", "CCCCCC", 4)
    _add_run(p,
        "Submitted by Oxydata Software Sdn Bhd  |  parikshit@oxydata.my"
        "  |  +603-7625 8298  |  www.oxydata.my",
        size_pt=8.5, color="gray", italic=True,
    )

    p = doc.add_paragraph()
    _set_para_spacing(p, 20, 0)
    _add_run(p,
        "Confidential — for evaluation purposes only."
        " Do not contact candidate directly without consent.",
        size_pt=8.5, color="gray", italic=True,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# LLM: infer candidate metadata from CV text
# ══════════════════════════════════════════════════════════════════════════════

def infer_candidate_metadata(oa: OpenAI, resume_text: str, job_name: str) -> Dict[str, str]:
    prompt = f"""You are a recruitment assistant. Read the CV below and extract the following fields.
Return ONLY a JSON object with exactly these keys:
- "location": candidate's current city and country (e.g. "Kuala Lumpur, Malaysia"). If not found, use "[Location — please verify]"
- "nationality": candidate's nationality (e.g. "Malaysian"). If not found, use "[Nationality — please verify]"
- "relevant_experience": years of experience directly relevant to the role of {job_name} (e.g. "10 Years"). Count only directly relevant years. If unclear, use "[Experience — please verify]"

Return only valid JSON. No explanation, no markdown, no backticks.

CV:
{resume_text[:6000]}"""

    try:
        r = oa.chat.completions.create(
            model=Config.OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "Return only valid JSON. No markdown, no explanation."},
                {"role": "user",   "content": prompt},
            ],
            temperature=0.1,
            max_tokens=200,
        )
        text = (r.choices[0].message.content or "{}").strip()
        text = text.replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except Exception as e:
        print(f"  [WARN] Metadata inference failed: {e}")
        return _PLACEHOLDER_META.copy()


_PLACEHOLDER_META: Dict[str, str] = {
    "location":            "[Location — please verify]",
    "nationality":         "[Nationality — please verify]",
    "relevant_experience": "[Experience — please verify]",
}


# ══════════════════════════════════════════════════════════════════════════════
# Main batch loop
# ══════════════════════════════════════════════════════════════════════════════

def main() -> int:
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("job_id", help="Manatal Job ID (numeric)")
    parser.add_argument("--force", action="store_true",
                        help="Regenerate .docx for all qualifying candidates, "
                             "even if traffic_rpt already exists (use after rubric change)")
    args = parser.parse_args()

    if not args.job_id.isdigit():
        print(f"ERROR: job_id must be numeric, got: {args.job_id}", file=sys.stderr)
        return 2

    job_id = int(args.job_id)
    force  = args.force

    try:
        Config.validate()
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 2

    at = AirtableClient()
    threshold = Config.PASS_THRESHOLD

    print(f"\n{'='*70}")
    print(f"SUBMISSION REPORT GENERATOR (.docx)")
    print(f"{'='*70}")
    print(f"Job ID:    {job_id}")
    print(f"Threshold: {threshold}")
    if force:
        print(f"Mode:      FORCE — existing traffic_rpt will be overwritten")
    print(f"{'='*70}\n")

    # Load rubric for soft_skill category split
    print("Loading rubric from Airtable...")
    rubric = at.get_rubric(str(job_id))
    if not rubric:
        print("[WARN] No rubric found — category split will be empty")
        rubric = {}

    soft_skill_ids: set = set()
    for item in rubric.get("requirements", {}).get("must_have", []):
        if item.get("category", "").lower() == "soft_skill":
            soft_skill_ids.add(item.get("id", ""))

    # Fetch candidate records
    print(f"Fetching Airtable records for job_id={job_id}...")
    records = at.get_records_by_formula(f"{{job_id}}={job_id}")
    print(f"Found {len(records)} total record(s)\n")

    qualifying = [
        rec for rec in records
        if float(rec["fields"].get("t2_score") or rec["fields"].get("t1_score") or 0) >= threshold
    ]
    print(f"Candidates scoring >= {threshold}: {len(qualifying)}")

    oa = OpenAI(api_key=Config.OPENAI_API_KEY)
    generated = 0
    skipped   = 0

    for rec in qualifying:
        record_id = rec["id"]
        fields    = rec["fields"]
        full_name = fields.get("full_name", "Candidate")
        job_name  = fields.get("job_name",  f"Job {job_id}")

        if fields.get("traffic_rpt") and not force:
            print(f"  Skipped (traffic_rpt exists): {full_name}")
            skipped += 1
            continue

        print(f"\nProcessing: {full_name}")

        overall_score = float(fields.get("t2_score") or fields.get("t1_score") or 0)
        recommendation = "PASS" if overall_score >= threshold else "FAIL"

        raw_json = fields.get("ai_detailed_json", "{}")
        try:
            detailed_json: Dict[str, Any] = json.loads(raw_json) if isinstance(raw_json, str) else (raw_json or {})
        except Exception:
            detailed_json = {}

        cv_text = (fields.get("cv_text") or "").strip()
        resume_text = cv_text if cv_text and "no resume attached" not in cv_text.lower() else ""

        if resume_text:
            print(f"  Inferring metadata from CV ({len(resume_text)} chars)...")
            metadata = infer_candidate_metadata(oa, resume_text, job_name)
        else:
            print(f"  No CV text — using placeholder metadata")
            metadata = _PLACEHOLDER_META.copy()

        all_must_have    = detailed_json.get("must_have",    [])
        compliance       = detailed_json.get("compliance",   [])
        nice_to_have     = detailed_json.get("nice_to_have", [])
        technical_items  = [i for i in all_must_have if i.get("id") not in soft_skill_ids]
        soft_skill_items = [i for i in all_must_have if i.get("id") in soft_skill_ids]

        payload = {
            "full_name":           full_name,
            "job_name":            job_name,
            "overall_score":       overall_score,
            "recommendation":      recommendation,
            "ai_summary":          fields.get("ai_summary", ""),
            "compliance":          compliance,
            "technical":           technical_items,
            "soft_skill":          soft_skill_items,
            "nice_to_have":        nice_to_have,
            "location":            metadata.get("location",            _PLACEHOLDER_META["location"]),
            "nationality":         metadata.get("nationality",         _PLACEHOLDER_META["nationality"]),
            "relevant_experience": metadata.get("relevant_experience", _PLACEHOLDER_META["relevant_experience"]),
            "report_date":         datetime.now().strftime("%d %B %Y"),
            "show_risk":           False,
        }

        print(f"  Building .docx...")
        try:
            docx_bytes = build_docx(payload)
        except Exception as e:
            print(f"  [ERROR] .docx build failed: {e}")
            continue

        print(f"  .docx built ({len(docx_bytes):,} bytes)")

        # Clear existing attachment before uploading
        try:
            at.update_record(record_id, {"traffic_rpt": []})
        except Exception as e:
            print(f"  [WARN] Could not clear traffic_rpt: {e}")

        filename = safe_filename(f"{full_name}_{job_name}_submission") + ".docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        result = at.upload_attachment_from_bytes(
            record_id=record_id,
            field_name="traffic_rpt",
            file_content=docx_bytes,
            filename=filename,
            content_type=content_type,
        )
        if result:
            print(f"  Uploaded to Airtable: {filename}")
            generated += 1
        else:
            print(f"  [ERROR] Upload failed for {full_name}")

    print(f"\n{'='*70}")
    print(f"Submission Report Generation Complete")
    print(f"{'='*70}")
    print(f"  Generated : {generated}")
    print(f"  Skipped   : {skipped}")
    print(f"{'='*70}\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
