"""Convert SOP.md into a formatted Word document (SOP.docx)."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
MD_PATH = HERE / "SOP.md"
OUT_PATH = HERE / "SOP.docx"
SCREENSHOT_DIR = HERE / "screenshots"

IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
CHECKBOX_RE = re.compile(r"^- \[ \] (.+)$")


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_table(doc, header_cells, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, text in enumerate(header_cells):
        cell = hdr.cells[i]
        cell.text = text.strip()
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2D3748")

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = text.strip()
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    return table


def process_inline(paragraph, text):
    """Add text to a paragraph, handling **bold** and `code` spans."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x46, 0xC1)
        else:
            paragraph.add_run(part)


def build_docx():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    for hs in ["Heading 1", "Heading 2", "Heading 3"]:
        s = doc.styles[hs]
        s.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    i = 0
    table_buf = None  # (header_cells, rows)

    def flush_table():
        nonlocal table_buf
        if table_buf:
            hdr, rows = table_buf
            add_table(doc, hdr, rows)
            doc.add_paragraph()
            table_buf = None

    while i < len(lines):
        line = lines[i]

        # Skip markdown horizontal rules
        if line.strip() in ("---", "***", "___"):
            flush_table()
            i += 1
            continue

        # Table row
        m_table = TABLE_ROW_RE.match(line.strip())
        if m_table:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Separator row (e.g., |---|---|)
            if all(re.match(r"^-+$", c.strip()) for c in cells):
                i += 1
                continue
            if table_buf is None:
                table_buf = (cells, [])
            else:
                table_buf[1].append(cells)
            i += 1
            continue
        else:
            flush_table()

        # Image
        m_img = IMG_RE.match(line.strip())
        if m_img:
            alt, src = m_img.group(1), m_img.group(2)
            img_path = HERE / src
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(alt)
                run.italic = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
            i += 1
            continue

        # HTML comment (skip)
        if line.strip().startswith("<!--"):
            i += 1
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Checkbox list items
        m_cb = CHECKBOX_RE.match(line.strip())
        if m_cb:
            p = doc.add_paragraph(style="List Bullet")
            process_inline(p, m_cb.group(1))
            i += 1
            continue

        # Unordered list
        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            process_inline(p, line.strip()[2:])
            i += 1
            continue

        # Ordered list
        m_ol = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            process_inline(p, m_ol.group(2))
            i += 1
            continue

        # Indented sub-list (4 spaces + -)
        if re.match(r"^\s{4}- ", line):
            p = doc.add_paragraph(style="List Bullet 2")
            process_inline(p, line.strip()[2:])
            i += 1
            continue

        # Bold-only lines (used for sub-headers like "What success looks like:")
        if line.strip().startswith("**") and line.strip().endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip()[2:-2])
            run.bold = True
            run.font.size = Pt(10.5)
            i += 1
            continue

        # Italic-only lines (e.g., *End of SOP*)
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip().strip("*"))
            run.italic = True
            i += 1
            continue

        # Regular paragraph
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph()
            process_inline(p, stripped)

        i += 1

    flush_table()
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build_docx()
